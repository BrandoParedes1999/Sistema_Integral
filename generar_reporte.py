from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
AZUL      = RGBColor(0, 61, 165)   # #003da5
AZUL_CLARO= RGBColor(220, 230, 250)
GRIS      = RGBColor(80, 80, 80)
VERDE     = RGBColor(16, 185, 129)
NEGRO     = RGBColor(0, 0, 0)
BLANCO    = RGBColor(255, 255, 255)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=AZUL):
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        # línea inferior
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '003da5')
        pBdr.append(bot)
        pPr.append(pBdr)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, bold_prefix=None, bullet=False):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bullet:
        p.style = doc.styles['List Bullet']
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.color.rgb = AZUL
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = GRIS
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # cabecera
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, '003da5')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(10)
    # filas
    for ri, row in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'EEF2FF'
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = NEGRO
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("INFOMATRIX — DEFENSA DE PROYECTO")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Categoría: Desarrollo de Software")
r.font.size = Pt(14); r.font.color.rgb = GRIS

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("─" * 55)
r.font.color.rgb = AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UniSalud")
r.bold = True; r.font.size = Pt(32); r.font.color.rgb = AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema Integral de Salud Universitaria — UNACAR")
r.font.size = Pt(13); r.font.color.rgb = GRIS; r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run(f"Fecha: {datetime.date.today().strftime('%d de %B de %Y')}")
r.font.size = Pt(11); r.font.color.rgb = GRIS

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCIÓN DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Descripción General del Proyecto")
add_body(doc,
    "UniSalud es una plataforma web de gestión integral de salud universitaria desarrollada para la "
    "Universidad Autónoma del Carmen (UNACAR). Digitaliza y centraliza el seguimiento de la salud "
    "física, mental y nutricional de los estudiantes, permitiendo a los administradores capturar datos, "
    "generar reportes automáticos y visualizar estadísticas institucionales en tiempo real.")

add_heading(doc, "Problema que resuelve", level=2)
add_body(doc,
    "UNACAR carecía de un sistema digital unificado para el monitoreo de salud estudiantil. "
    "Los registros eran físicos, no existían alertas para casos críticos de salud mental, ni "
    "análisis estadísticos para la toma de decisiones institucionales.")

add_heading(doc, "Propuesta de valor", level=2)
for item in [
    "Integración de salud MENTAL + FÍSICA + NUTRICIÓN en una sola plataforma.",
    "Alertas automáticas para estudiantes con niveles severos de depresión/ansiedad (DASS-21).",
    "Generación automática de reportes PDF personalizados y envío por correo electrónico.",
    "Observatorio estadístico con gráficas interactivas filtradas por facultad, carrera y género.",
    "Acceso seguro para alumnos mediante OTP (One-Time Password) por correo.",
    "Planificador nutricional personalizado basado en el Gasto Energético Total (GET) del alumno.",
]:
    add_body(doc, item, bullet=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. RÚBRICAS DE EVALUACIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Rúbricas de Evaluación INFOMATRIX — Desarrollo de Software")

add_heading(doc, "Rubro 1 — Exposición", level=2)
add_body(doc, "Criterios generales de todas las categorías:", bold_prefix="▶")
for item in [
    "Claridad: Explicar con lenguaje sencillo qué hace el sistema, para quién y por qué.",
    "Desenvolvimiento y apropiación: Demostrar dominio total del sistema; el jurado detecta quién lo hizo.",
    "Ejemplifica y demuestra: Hacer una demo EN VIVO del sistema (login, captura, reporte PDF, observatorio).",
    "Trabajo en equipo: Todos los integrantes deben poder responder preguntas de cualquier módulo.",
    "Demostración (específico Desarrollo de Software): Mostrar el software funcionando con datos reales.",
]:
    add_body(doc, item, bullet=True)

add_heading(doc, "Rubro 2 — Proyecto", level=2)
for item in [
    "Originalidad e Innovación: Sistema único que combina 3 dimensiones de salud con instrumentos clínicos validados.",
    "Comparativa con mercado: Mostrar tabla comparativa vs Google Forms, EHR genéricos y soluciones de pago.",
    "Propuesta de solución a un problema real de su contexto: Digitalización de la salud estudiantil en UNACAR.",
]:
    add_body(doc, item, bullet=True)

add_heading(doc, "Rubro 3 — Mensaje / Impacto", level=2)
for item in [
    "Mensaje: 'La salud universitaria no debe ser un archivo en el cajón.'",
    "Impacto en sociedad: Detección temprana de enfermedades crónicas y salud mental en jóvenes.",
    "UX: OTP sin contraseña, semáforo de colores, todo visible en una pantalla.",
    "UI: Diseño institucional UNACAR, Bootstrap 5, responsive, gráficas interactivas.",
    "Fuerza visual: Stand con capturas, diagrama de arquitectura y ejemplo de reporte PDF impreso.",
]:
    add_body(doc, item, bullet=True)

add_heading(doc, "Rubro 4 — Complejidad Técnica", level=2)
for item in [
    "15+ tablas relacionadas en MySQL (alumnos, DASS, datos físicos, menú, ingredientes, etc.).",
    "Sistema de bloqueo progresivo de cuentas (2min → 5min → 30min → 1hr → 24hr).",
    "Cálculo automático de GET con fórmula metabólica individualizada.",
    "Generación dinámica de PDFs con FPDF y envío automático vía PHPMailer/SMTP.",
    "Dos flujos de autenticación distintos: Admin (contraseña) y Alumno (OTP).",
    "Control de acceso por roles: Admin, Capturista, Alumno.",
    "Sanitización de inputs contra SQL Injection y XSS en todo el sistema.",
    "Código o Diagrama de Bloques debe estar impreso y disponible en el stand.",
]:
    add_body(doc, item, bullet=True)

add_heading(doc, "Rubro 5 — Materiales del Stand", level=2)
add_table(doc,
    ["Material", "¿Tienes?", "Prioridad"],
    [
        ["Laptop con el sistema corriendo (demo en vivo)", "✅ Confirmar", "🔴 Crítico"],
        ["Reporte del proyecto impreso", "Preparar", "🔴 Crítico"],
        ["Bitácora de desarrollo", "Preparar", "🔴 Crítico"],
        ["Diagrama de arquitectura / bloques impreso", "Preparar", "🔴 Crítico"],
        ["Tabla comparativa con mercado", "Preparar", "🟡 Importante"],
        ["Ejemplo de reporte PDF generado por el sistema", "Generar", "🟡 Importante"],
        ["Cartel/banner con nombre y propuesta de valor", "Opcional", "🟢 Recomendado"],
        ["QR hacia demo o video del sistema", "Opcional", "🟢 Recomendado"],
    ],
    [7, 3, 3]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. METODOLOGÍA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Metodologías Aplicadas en el Desarrollo")

add_heading(doc, "3.1 Metodología de Desarrollo: SCRUM (Ágil)", level=2)
add_body(doc,
    "Se adoptó un enfoque ágil basado en SCRUM para gestionar el desarrollo iterativo e incremental "
    "del sistema. Esto permitió entregar módulos funcionales en ciclos cortos (sprints) y adaptar "
    "los requerimientos conforme avanzaba el proyecto.")
for item in [
    "Sprint 1: Análisis de requerimientos, diseño de base de datos y módulo de autenticación.",
    "Sprint 2: Módulo de gestión de alumnos, captura de datos físicos y generación de PDFs.",
    "Sprint 3: Cuestionarios DASS-21 y PEPS-1 con lógica de puntuación.",
    "Sprint 4: Observatorio estadístico con gráficas interactivas.",
    "Sprint 5: Planificador nutricional y módulo de credenciales digitales.",
    "Sprint 6: Pruebas de integración, corrección de errores, despliegue.",
]:
    add_body(doc, item, bullet=True)

add_heading(doc, "3.2 Metodología de Diseño: Design Thinking", level=2)
add_body(doc,
    "Para definir las necesidades reales de los usuarios (administradores, capturistas y alumnos) "
    "se aplicó Design Thinking en 5 fases:")
for item in [
    "Empatizar: Entrevistas con personal de salud de UNACAR para entender sus procesos actuales.",
    "Definir: El problema central era la falta de trazabilidad y análisis de datos de salud estudiantil.",
    "Idear: Propuesta de plataforma web con módulos integrados y roles diferenciados.",
    "Prototipar: Wireframes de la UI para validar flujos antes de programar.",
    "Evaluar: Pruebas con usuarios reales para ajustar la experiencia de uso.",
]:
    add_body(doc, item, bullet=True)

add_heading(doc, "3.3 Metodología Científica en los Instrumentos de Salud", level=2)
add_body(doc,
    "Los cuestionarios integrados en UniSalud están basados en instrumentos clínicos validados internacionalmente:")
for item in [
    "DASS-21 (Depression Anxiety Stress Scales): Instrumento psicométrico con 21 ítems validado para medir depresión, ansiedad y estrés en población universitaria. Escala de 0-3 por ítem, con clasificaciones clínicas establecidas.",
    "PEPS-1 (Estilo de Vida): Cuestionario de 48 ítems para evaluar hábitos de salud y comportamientos de riesgo en estudiantes universitarios.",
    "Cálculo de GET: Basado en fórmulas metabólicas reconocidas (Harris-Benedict o Mifflin-St Jeor) ajustadas por nivel de actividad física.",
]:
    add_body(doc, item, bullet=True)

add_heading(doc, "3.4 Arquitectura del Sistema: MVC (Modelo-Vista-Controlador)", level=2)
add_body(doc,
    "Aunque el proyecto usa PHP procedural/orientado a objetos, sigue los principios del patrón MVC:")
add_table(doc,
    ["Capa", "Componente en UniSalud", "Tecnología"],
    [
        ["Modelo (Model)", "Base de datos, clases de conexión, lógica de negocio", "MySQL + PHP (config/config.php)"],
        ["Vista (View)", "Interfaces HTML/CSS, dashboards, formularios", "HTML5 + Bootstrap 5 + JS"],
        ["Controlador (Controller)", "Endpoints PHP que reciben, procesan y responden peticiones", "PHP (guardar_*.php, obtener_*.php)"],
    ],
    [4, 7, 6]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. COSTO DEL PROYECTO Y JUSTIFICACIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Costo del Proyecto y Justificación de Precio")

add_heading(doc, "4.1 Desglose de Costos de Desarrollo", level=2)
add_body(doc,
    "El cálculo está basado en tarifas del mercado mexicano para desarrolladores web "
    "(junior-semi senior), estimando las horas reales de trabajo por módulo.")

add_table(doc,
    ["Módulo / Actividad", "Horas Est.", "Tarifa/hr (MXN)", "Subtotal (MXN)"],
    [
        ["Análisis y diseño de base de datos (15+ tablas)", "20 hrs", "$200", "$4,000"],
        ["Sistema de autenticación dual (Admin OTP + Alumno)", "25 hrs", "$200", "$5,000"],
        ["Módulo gestión de alumnos (CRUD completo)", "30 hrs", "$200", "$6,000"],
        ["Captura de datos físicos + generación PDF (FPDF)", "35 hrs", "$250", "$8,750"],
        ["Cuestionario DASS-21 con lógica de puntuación", "20 hrs", "$250", "$5,000"],
        ["Cuestionario PEPS-1 (48 ítems + resultados)", "15 hrs", "$250", "$3,750"],
        ["Observatorio estadístico (5 métricas + gráficas)", "40 hrs", "$250", "$10,000"],
        ["Planificador nutricional semanal (JS + PHP)", "45 hrs", "$300", "$13,500"],
        ["Sistema de credenciales digitales PDF", "20 hrs", "$200", "$4,000"],
        ["Reportes y exportación Excel/PDF", "20 hrs", "$200", "$4,000"],
        ["Portal del alumno (dashboard, perfil, historial)", "30 hrs", "$250", "$7,500"],
        ["Seguridad (bloqueo cuentas, tokens, sanitización)", "15 hrs", "$300", "$4,500"],
        ["Integración PHPMailer + SMTP corporativo", "10 hrs", "$200", "$2,000"],
        ["Diseño UI/UX (CSS, responsive, colores UNACAR)", "25 hrs", "$200", "$5,000"],
        ["Pruebas, depuración y documentación", "30 hrs", "$200", "$6,000"],
        ["TOTAL DE DESARROLLO", "380 hrs", "—", "$89,000 MXN"],
    ],
    [7, 2.5, 3.5, 3.5]
)

add_heading(doc, "4.2 Costos de Infraestructura (Anuales)", level=2)
add_table(doc,
    ["Concepto", "Costo Anual (MXN)", "Notas"],
    [
        ["Hospedaje web (hosting cPanel compartido)", "$1,800", "p. ej. Hostinger, SiteGround"],
        ["Dominio (.com.mx)", "$350", "Renovación anual"],
        ["Correo corporativo (Neubox / Google Workspace)", "$1,200", "Actual: Neubox configurado"],
        ["SSL/TLS Certificate (HTTPS)", "$0", "Let's Encrypt gratuito"],
        ["Backups automáticos", "$600", "Incluido en hosting premium"],
        ["TOTAL INFRAESTRUCTURA ANUAL", "$3,950 MXN/año", "≈ $329/mes"],
    ],
    [7, 4, 6]
)

add_heading(doc, "4.3 Costo Total del Proyecto", level=2)
add_table(doc,
    ["Concepto", "Monto (MXN)", "Monto (USD aprox.)"],
    [
        ["Desarrollo del sistema (una sola vez)", "$89,000", "≈ $4,750 USD"],
        ["Infraestructura primer año", "$3,950", "≈ $210 USD"],
        ["TOTAL AÑO 1", "$92,950 MXN", "≈ $4,960 USD"],
        ["Mantenimiento anual (desde año 2)", "$15,000 + $3,950", "≈ $1,010 USD/año"],
    ],
    [8, 4, 5]
)

add_heading(doc, "4.4 Justificación del Precio", level=2)
add_body(doc, "¿Por qué vale $89,000 MXN el desarrollo?", bold_prefix="▶")
for item in [
    "380 horas de desarrollo real con múltiples tecnologías integradas (PHP, MySQL, JavaScript, FPDF, PHPMailer, Bootstrap).",
    "Instrumentos clínicos validados (DASS-21, PEPS-1) implementados con lógica de puntuación especializada, no es un formulario genérico.",
    "Sistema de seguridad robusto: bloqueo progresivo, tokens internos, sanitización completa, roles diferenciados.",
    "Generación automática de PDFs personalizados por alumno con datos clínicos — funcionalidad que en EHR comerciales cuesta $$$.",
    "Observatorio estadístico con 5+ métricas de salud, filtros dinámicos y visualizaciones — equivale a un módulo de Business Intelligence.",
    "Planificador nutricional con base de datos de platillos, ingredientes y grupos alimenticios — es un submódulo completo.",
    "Comparativa: Un EHR (Electronic Health Record) comercial como MedScape, Nuvolo o eClinicalWorks cobra entre $150-$500 USD/mes por institución.",
    "UniSalud fue desarrollado a medida para el contexto universitario mexicano, en español y sin costos de licencia perpetua.",
]:
    add_body(doc, item, bullet=True)

add_heading(doc, "4.5 Comparativa de Costo vs. Soluciones Comerciales", level=2)
add_table(doc,
    ["Solución", "Costo Mensual", "Costo 5 Años", "Personalización UNACAR", "En Español"],
    [
        ["UniSalud (este proyecto)", "$329 MXN hosting", "~$107,000 MXN", "100% ✅", "✅"],
        ["Google Forms + Sheets", "$0", "$0", "Muy limitada ❌", "Parcial"],
        ["EHR Comercial (eClinicalWorks)", "~$9,000 MXN/mes", "~$540,000 MXN", "Mínima ❌", "Parcial"],
        ["Sistema a medida (agencia)", "N/A", "~$200,000+ MXN", "✅", "✅"],
        ["Microsoft Power Apps + SharePoint", "~$4,500 MXN/mes", "~$270,000 MXN", "Media", "✅"],
    ],
    [5.5, 3.5, 3.5, 4, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PREGUNTAS FRECUENTES DEL JURADO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Preguntas Frecuentes del Jurado y Respuestas Recomendadas")

preguntas = [
    (
        "¿Por qué PHP y no Node.js, Python o Java?",
        "PHP es ampliamente usado en sistemas web de instituciones educativas en México, tiene despliegue nativo en servidores Apache sin configuración adicional, no requiere licencias, y tiene una curva de aprendizaje accesible. Para el alcance de este proyecto, PHP fue la opción más eficiente y práctica. Además, librerías clave como FPDF y PHPMailer tienen soporte nativo en PHP."
    ),
    (
        "¿Cómo protegen los datos clínicos de los alumnos?",
        "El sistema implementa: (1) Autenticación por roles — ningún usuario accede a datos que no le corresponden, (2) Sanitización de todos los inputs con htmlspecialchars() y prepared statements para prevenir SQL Injection y XSS, (3) Sistema de bloqueo progresivo de cuentas, (4) Tokens internos para operaciones de fondo (generación de PDFs), (5) Sesiones PHP con verificación en cada página protegida."
    ),
    (
        "¿El DASS-21 es un instrumento médico válido?",
        "Sí. El DASS-21 (Depression Anxiety Stress Scales) es un instrumento psicométrico desarrollado por Lovibond & Lovibond (1995) con amplia validación científica internacional. Tiene versiones validadas para población universitaria latinoamericana. No reemplaza un diagnóstico clínico, pero es una herramienta de tamizaje ampliamente utilizada en departamentos de salud universitaria."
    ),
    (
        "¿Ya está implementado en producción en UNACAR?",
        "El sistema está desarrollado y listo para implementarse. [Si ya hay implementación: mencionar cuántos alumnos, capturas, reportes generados]. La base de datos incluye el esquema completo y datos de prueba. El siguiente paso es el despliegue en el servidor institucional de UNACAR y la capacitación del personal de salud."
    ),
    (
        "¿Qué diferencia hay con simplemente usar Excel o Google Forms?",
        "Excel y Google Forms son herramientas genéricas sin lógica clínica. UniSalud: calcula automáticamente el IMC, la clasificación DASS-21, el GET nutricional y genera alertas; produce reportes PDF personalizados y los envía por correo sin intervención manual; consolida datos de múltiples fuentes en un observatorio estadístico; tiene control de acceso por roles; y mantiene historial clínico por alumno a lo largo del tiempo."
    ),
    (
        "¿Qué tan escalable es el sistema?",
        "La arquitectura modular permite agregar nuevos cuestionarios, métricas o módulos sin modificar los existentes. La base de datos está normalizada. El sistema podría escalar a otras universidades con mínimas adaptaciones de configuración (institución, colores, SMTP). Con un servidor más robusto (VPS/cloud) soportaría miles de usuarios concurrentes."
    ),
    (
        "¿Cómo calculan el GET (Gasto Energético Total)?",
        "Usando la fórmula de Harris-Benedict ajustada por factor de actividad física: se toma el peso, talla, edad y sexo del alumno para calcular la Tasa Metabólica Basal (TMB), luego se multiplica por el factor de actividad (sedentario 1.2, ligero 1.375, moderado 1.55, activo 1.725). Este resultado determina las kilocalorías diarias recomendadas y la distribución de macronutrientes (proteínas 20%, carbohidratos 55%, grasas 25%)."
    ),
    (
        "¿Tienen patente o derechos de autor sobre el software?",
        "El software fue desarrollado por nosotros y los derechos pertenecen a los autores. Para protección formal, el código puede registrarse ante el INDAUTOR (Instituto Nacional del Derecho de Autor) de México como obra literaria de software. También es posible documentarlo como Modelo de Utilidad ante el IMPI si se implementa una solución tecnológica novedosa."
    ),
    (
        "¿Cuánto tiempo tardaron en desarrollarlo?",
        "El desarrollo total implicó aproximadamente 380 horas de trabajo distribuidas en sprints bajo metodología ágil SCRUM. Se trabajó desde el análisis de requerimientos con el personal de salud de UNACAR, pasando por el diseño de la base de datos, hasta las pruebas de integración y ajustes de UX con usuarios reales."
    ),
    (
        "¿Qué pasaría si un alumno tiene un resultado muy alto en depresión?",
        "El sistema está diseñado para generar alertas visibles en el dashboard del administrador cuando un alumno obtiene puntuaciones en nivel Severo o Extremadamente Severo en cualquiera de las dimensiones del DASS-21 (Depresión ≥28, Ansiedad ≥20, Estrés ≥34). Esto permite al personal de salud de UNACAR contactar proactivamente al estudiante. El sistema muestra también información sobre la clínica universitaria dentro del perfil del alumno."
    ),
    (
        "¿Cómo generan las credenciales digitales?",
        "El sistema utiliza la librería FPDF para generar PDFs con los datos del alumno (nombre, matrícula, foto, facultad, carrera). Se puede generar de forma individual o en lote para todos los alumnos. Los PDFs incluyen un código QR que enlaza al perfil del alumno. El diseño sigue la identidad gráfica institucional de UNACAR."
    ),
    (
        "¿Qué impacto social real tiene el proyecto?",
        "1) Detección temprana de salud mental: permite identificar estudiantes en riesgo antes de que lleguen a crisis. 2) Datos para políticas institucionales: el observatorio permite que UNACAR tome decisiones informadas sobre programas de salud. 3) Empoderamiento del alumno: cada estudiante puede ver su historial de salud, descargarlo y tomar decisiones informadas. 4) Reducción de costos: elimina registros físicos, optimiza el tiempo del personal de salud y evita la contratación de sistemas costosos."
    ),
]

for i, (pregunta, respuesta) in enumerate(preguntas, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f"P{i}: {pregunta}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = AZUL

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.7)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run("R: " + respuesta)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRIS

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. GUIÓN DE DEFENSA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Guión Recomendado para la Defensa (5-7 minutos)")

pasos = [
    ("0:00 - 0:30", "APERTURA",
     "Presentarse, mencionar la institución (UNACAR) y lanzar la pregunta gancho: "
     "'¿Cuántos datos de salud de los estudiantes universitarios se pierden cada semestre en papel?' "
     "Presentar UniSalud en una oración."),
    ("0:30 - 1:30", "EL PROBLEMA",
     "Explicar la situación actual en UNACAR: registros físicos, sin análisis, sin alertas para casos "
     "críticos de salud mental, sin seguimiento nutricional. Mostrar el impacto en números si es posible "
     "(cuántos alumnos, cuántas capturas al semestre)."),
    ("1:30 - 4:30", "DEMO EN VIVO",
     "1) Login admin → Dashboard con estadísticas. "
     "2) Captura de datos físicos de un alumno → PDF generado automáticamente. "
     "3) Resultado DASS-21 con clasificación y alerta. "
     "4) Planificador nutricional semanal. "
     "5) Observatorio con gráfica de IMC por facultad."),
    ("4:30 - 5:30", "DIFERENCIADORES",
     "Mostrar brevemente la tabla comparativa. Destacar: costo cero de licencias, hecho a medida para UNACAR, "
     "instrumentos clínicos validados, seguridad implementada."),
    ("5:30 - 6:00", "IMPACTO Y CIERRE",
     "Mencionar el impacto social: detección temprana, datos para decisiones institucionales, "
     "empoderamiento del alumno. Cerrar con: 'UniSalud convierte datos de salud en decisiones informadas.'"),
    ("6:00+", "PREGUNTAS",
     "Escuchar con calma. Si no saben algo, ser honestos: 'Ese aspecto está planeado para la siguiente versión' "
     "o 'Lo investigamos y les confirmamos'. No improvisar datos médicos o técnicos sin estar seguros."),
]

for tiempo, titulo, desc in pasos:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"[{tiempo}]  {titulo}")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
    add_body(doc, desc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. STACK TECNOLÓGICO RESUMIDO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "7. Stack Tecnológico")
add_table(doc,
    ["Capa", "Tecnología", "Versión", "Función"],
    [
        ["Backend", "PHP", "7.x+", "Lógica de negocio, endpoints, sesiones"],
        ["Base de datos", "MySQL / MariaDB", "5.7+", "Almacenamiento, 15+ tablas relacionadas"],
        ["PDF", "FPDF", "1.86", "Generación de reportes y credenciales"],
        ["Email", "PHPMailer", "7.0.0", "Envío de OTP y reportes por SMTP"],
        ["Frontend", "Bootstrap", "5.3.2", "UI responsive, grid, componentes"],
        ["Frontend", "JavaScript Vanilla", "ES6+", "Planificador, validaciones, AJAX"],
        ["Iconos", "Bootstrap Icons", "1.11.3", "Iconografía del sistema"],
        ["Servidor", "Apache + .htaccess", "2.4+", "Servidor web, rutas, seguridad"],
        ["Seguridad", "PHP Sessions + Tokens", "—", "Autenticación, control de acceso"],
    ],
    [3, 3.5, 2.5, 8]
)

# ── Pie de página ─────────────────────────────────────────────────────────────
section = doc.sections[0]
footer  = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(f"UniSalud — Reporte de Defensa INFOMATRIX  |  {datetime.date.today().year}  |  Categoría: Desarrollo de Software")
fr.font.size = Pt(8)
fr.font.color.rgb = GRIS

path = "/home/user/Sistema_Integral/UniSalud_Defensa_INFOMATRIX.docx"
doc.save(path)
print(f"Documento guardado: {path}")
